import os
import os.path
import pathlib
import time
import functools
import sqlite3
import io
from PIL import Image

import flask
from flask import Flask
from flask import request
from flask import Blueprint, g, current_app, session
from werkzeug.middleware.proxy_fix import ProxyFix
import werkzeug.utils
import logging
import click
import htmldocx
import docx

from . import document
from . import storage


def create_app(instance_path=None):
  if instance_path is None:
    app = Flask(__name__, instance_relative_config=True)
  else:
    app = Flask(__name__, instance_relative_config=True,
                instance_path=instance_path)
  app.config.from_mapping(
    SECRET_KEY='DEV',
    DATABASE=os.path.join(app.instance_path, 'docserve.sqlite'),    
  )
  app.config.from_prefixed_env()
  app.config.from_pyfile('config.py', silent=True)

  # Configure logging
  BASE_DIR = os.getcwd()
  log_file_name = os.path.join(BASE_DIR, 'log-docserve.txt')
  FORMAT = '%(asctime)s:%(levelname)s:%(name)s:%(message)s'
  logging.basicConfig(filename=log_file_name,
                      level=logging.INFO,
                      format=FORMAT,)
  
  logging.info("Starting docserve.server: %s", __name__)
  
  # If so configured, setup for running behind a reverse proxy.
  if app.config.get('PROXY_CONFIG'):
    app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1,
                            x_host=1, x_prefix=1)
    logging.info("set proxy fix")

  try:
    logging.info("instance path = %s", app.instance_path)
    os.makedirs(app.instance_path)
  except OSError as err:
    pass

  try:
    data_path = os.path.join(app.instance_path, 'data')
    logging.info("data path = %s", data_path)
    os.makedirs(data_path)
  except OSError as err:
    pass

  app.register_blueprint(bp)

  @app.errorhandler(Exception)
  def handle_exception(e):
    logging.exception("Internal error")
    return e
  return app



def get_db():
  if 'db' not in g:
    g.db = sqlite3.connect(
      current_app.config['DATABASE'],
      detect_types=sqlite3.PARSE_DECLTYPES)
    g.db.row_factory = sqlite3.Row
  return g.db

def close_db(e=None):
  db = g.pop('db', None)
  if db is not None:
    db.close()


def init_db():
  db = get_db()
  with current_app.open_resource('schema.sql') as f:
    db.executescript(f.read().decode('utf8'))


bp = Blueprint('docserve', __name__, cli_group=None)

@bp.cli.command('init-db')
def init_db_command():
  """Drop and recreate tables."""
  init_db()
  click.echo('Initialized database.')

@bp.cli.command('dump-db')
def init_db_command():
  """Dump table contents"""
  storage.dump_tables(get_db())

@bp.cli.command('xfer-user')
@click.argument('old_user_id')
@click.argument('new_user_id')
def xfer_user_id(old_user_id, new_user_id):
  """Replace old user with new user"""
  storage.xfer_user_id(get_db(), old_user_id, new_user_id)

@bp.route('/Privacy')
def serve_privacy():
  return flask.send_from_directory(current_app.root_path,
                                   'static/privacy')

def extract_auth_key(headers):
  auth = request.headers.get('Authorization')
  if auth is not None:
    index = auth.find(' ')
    if index > 0:
      return auth[index+1:]
  return None

def login_required(view):
  @functools.wraps(view)
  def wrapped_view(**kwargs):
    if current_app.config.get('AUTH_KEY'):
      if storage.check_address_block(get_db(), request.remote_addr):
        return { "error": "Denied" }, 401
      
      # Verify auth matches
      auth = extract_auth_key(request.headers)
      if auth != current_app.config['AUTH_KEY']:
        storage.inc_bad_access(get_db(), request.remote_addr)
        return { "error": "Invalid authorization header" }, 401

    return view(**kwargs)
  return wrapped_view

def user_id_required(view):
  @functools.wraps(view)
  def wrapped_view(**kwargs):
    if storage.check_address_block(get_db(), request.remote_addr):
      return "Denied", 401

    user_id = request.args.get("q", "")

    if not storage.check_user_id(get_db(), user_id):
      storage.inc_bad_access(get_db(), request.remote_addr)
      return "No auth", 401

    return view(**kwargs)
  return wrapped_view

def get_user_id(request):
  GPT_USER="Openai-Ephemeral-User-Id"
  GPT_CONVO="Openai-Conversation-Id"
  gpt_convo_id = request.headers.get(GPT_CONVO, "<blank convo>")
  gpt_user_id = request.headers.get(GPT_USER, "<blank user>")
  user_id = storage.get_user_id(get_db(), gpt_convo_id, gpt_user_id)
  logging.info("user_id = %s", user_id)
  return user_id
  

@bp.route('/docs', methods=["POST","GET"])
@login_required
def docs():
  """
  List and create docs
  """
  user_id = get_user_id(request)

  if request.method == "GET":
    # Get a list of docs
    doc_list = []
    content = { "documents": doc_list }
    doc_ids = storage.get_user_doc_ids(get_db(), user_id)
    for doc_id in doc_ids:
      logging.info(f"loading doc_id: {doc_id}")
      filename = storage.get_doc_filename(get_db(), user_id, doc_id)        
      doc = document.Document.loadDocument(storage.data_path(current_app),
                                           filename)
      title = doc.getElementContent(document.ElementType.TITLE, 0)  
      doc_list.append({ "doc_id": doc_id, "title": title})
    return flask.jsonify(content)      
  else:  # POST  
    # Create a new document
    logging.info("CreateDoc")
    content = request.json
    value = None
    if content is not None:
      value = content.get('title')
      doc_id = content.get('doc_id')
      logging.info("create doc id: %s, title: %s", doc_id, value)
    
    if value is None:
      return { "error": "Malformed input" }, 400

    # Check if doc_id already exists, assign another if so
    if (doc_id is not None and
        storage.get_doc_filename(get_db(), user_id, doc_id) is not None):
      doc_id = None  

    doc = document.Document.newDocument(doc_id)
    element_id = doc.setElementContent(document.ElementType.TITLE, 0, value)
    filename = storage.gen_filename()
    storage.add_user_doc(get_db(), user_id, doc.doc_id, filename)
    doc.saveDocument(storage.data_path(current_app), filename)  

    content = { "doc_id" : doc.docId(),
                "doc_view_url": flask.url_for('docserve.view_doc',
                                              _external=True),
                "image_upload_url": flask.url_for('docserve.upload_view',
                                                  _external=True),
                "user_id": user_id
               }
    return flask.jsonify(content)


@bp.route('/docs/<doc_id>')
@login_required
def read_doc(doc_id):
  """
  Read the details of a docuemnt
  """
  user_id = get_user_id(request)  
  logging.info("ReadDoc[%s]", doc_id)
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)

  if doc is None:
    return { "error": "Document not found" }, 400
  

  # Assemble doc properties
  content = { "doc_id" : doc_id,
              "doc_view_url": flask.url_for('docserve.view_doc',
                                            _external=True),
              "image_upload_url": flask.url_for('docserve.upload_view',
                                                _external=True),
              "user_id": user_id              
             }

  entries = []
  entries.append({ "index" : 0, "type" : document.ElementType.TITLE })
  entries.append({ "index" : 0, "type" : document.ElementType.ABSTRACT })
  entries.append({ "index" : 0, "type" : document.ElementType.OUTLINE })
  
  for element in doc.elements:
    entry = { "index" : element.elementIndex, "type" : element.elementType }
    entries.append(entry)
  content["elements"] = entries
  return flask.jsonify(content)


@bp.route('/docs/<doc_id>/elements/<type>', methods=["POST"])
@login_required
def create_element(doc_id, type):
  user_id = get_user_id(request)  
  element_type = type
  content = request.json
  logging.info("CreateElement[%s:%s]", doc_id, element_type)

  value = ""
  if content is not None:
    value = content.get('value', "")
  logging.info("CreateElement - value = %s", value)
    
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)
  if doc is None:
    return { "error": "Document not found" }, 400  

  element_index = doc.setElementContent(element_type, 0, value)
  doc.saveDocument(storage.data_path(current_app))

  content = { "doc_id" : doc.docId(),
              "type" : element_type,
              "index": element_index,
             }
  return flask.jsonify(content)


@bp.route('/docs/<doc_id>/elements/<type>/<index>', methods=["GET", "PUT"])
@login_required
def read_write_element(doc_id, type, index): 
  """
  Read and update document elements
  """
  user_id = get_user_id(request)  
  element_type = type  
  element_index = int(index)
  logging.info("Element[%s:%s:%d]", doc_id, element_type, element_index)
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)

  if doc is None:
    return { "error": "Document not found" }, 400  

  if request.method == "GET":
    # Read Element
    value = doc.getElementContent(element_type, element_index)
    logging.info("ReadElement result: %s", value)
  
    content = { "doc_id" : doc.docId(),
                "type" : element_type,
                "index": element_index,
                "value": value,
               }
    return flask.jsonify(content)

  else: # PUT
    value = None
    content = request.json    
    if content is not None:
      value = content.get('value')
    if value is None:
      return { "error": "Malformed input" }, 400      

    logging.info("WriteElement - value = %s", value)
    element_index = doc.setElementContent(element_type, element_index, value)
    doc.saveDocument(storage.data_path(current_app))

    content = { "doc_id" : doc.docId(),
                "type" : element_type,
                "index": element_index,
               }
    return flask.jsonify(content)



@bp.route("/list")
@user_id_required
def list_docs():
  user_id = request.args.get("q", "")
  doc_ids = storage.get_user_doc_ids(get_db(), user_id)
  entries = []  
  for doc_id in doc_ids:
    filename = storage.get_doc_filename(get_db(), user_id, doc_id)            
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)
    title = doc.getElementContent(document.ElementType.TITLE, 0)
    entries.append((doc.doc_id, title))
  return flask.render_template("list.html", entries=entries, user_id=user_id)

def image_url_func(user_id, doc_id, element_index, ref_val):
  url = flask.url_for('docserve.images', doc_id=doc_id,
                      element_index=element_index, q=user_id,
                      _external=True)
  logging.info("gen image url %s", url)
  return url

@bp.route("/view")
@user_id_required
def view_doc():
  user_id = request.args.get("q", "")
  doc_id = request.args.get("doc_id", "")  
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)
  if doc is None:
    return "Document not found", 400

  context = document.DocContext(image_url_func, user_id, doc_id)
  
  return flask.render_template("view.html", doc=doc,
                               user_id=user_id, context=context)


@bp.route("/images/<doc_id>/<element_index>", methods=["POST","GET"])
@user_id_required
def images(doc_id, element_index):
  logging.info("images %s", doc_id)  
  element_index = int(element_index)
  user_id = request.args.get("q", "")
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)
  if doc is None:
    return "Document not found", 400

  ref_val = doc.getElementContent(document.ElementType.IMAGE, element_index)
  if ref_val is None:
    logging.info("image element content not found: %d", element_index)
    return "Image not found", 400
  
  image_file = os.path.join(storage.data_path(current_app), filename +
                            "." + str(element_index) + ".png")
  logging.info("image file %s", image_file)

  if request.method == "GET":
    logging.info("get image")    
    if not os.path.isfile(image_file):
      return "Image file not found", 400
    return flask.send_file(image_file, mimetype="image/webp")

  else: # POST
    # Upload file
    logging.info("upload image")
    if ('file' not in request.files or
        request.files['file'].filename == ''):
      return flask.redirect(flask.url_for("docserve.upload_view",
                                          doc_id=doc_id,
                                          element_index=element_index,
                                          q=user_id))

    file = request.files['file']
    upload_filename = request.files['file'].filename
    image = Image.open(file).convert("RGB")
    logging.info("saving image: %s", image_file)
    image.save(image_file, "png")
    file.close()
    flask.flash(f"Uploaded image file: {upload_filename}")

    # Touch file to signal a change
    pathlib.Path(os.path.join(storage.data_path(current_app),
                              filename)).touch()

    return flask.redirect(flask.url_for("docserve.upload_view",
                                        doc_id=doc_id,
                                        element_index=element_index,
                                        q=user_id, done=True))
    


@bp.route('/upload')
@user_id_required
def upload_view():
  """
  Upload images
  """
  doc_id = request.args.get("doc_id", "")  
  element_index = int(request.args.get("index", "0"))
  user_id = request.args.get("q", "")
  done = request.args.get("done", False)
    
  return flask.render_template("upload.html", doc_id=doc_id,
                               user_id=user_id, element_index=element_index,
                               done=done)


@bp.route("/docx/<doc_id>")
@user_id_required
def download_docx(doc_id):
  user_id = request.args.get("q", "")
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  doc = None
  if filename is not None:
    doc = document.Document.loadDocument(storage.data_path(current_app),
                                         filename)
  if doc is None:
    return "Document not found", 400

  context = document.DocContext(image_url_func, user_id, doc_id)  

  parser = htmldocx.HtmlToDocx()
  expo_doc = docx.Document()
  parser.add_html_to_document(doc.titleElement.emitHTML(context), expo_doc)
  parser.add_html_to_document("<h2>Abstract</h2>", expo_doc)  
  parser.add_html_to_document(doc.abstractElement.emitHTML(context), expo_doc)
  parser.add_html_to_document("<h2>Outline</h2>", expo_doc) 
  parser.add_html_to_document(doc.outlineElement.emitHTML(context), expo_doc)
  logging.info("outline html: %s", doc.outlineElement.emitHTML(context))
  for element in doc.elements:
    parser.add_html_to_document(element.emitHTML(context), expo_doc)          
  out_file = io.BytesIO()
  expo_doc.save(out_file)
  out_file.seek(0, 0)
  mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
  return flask.send_file(out_file, mimetype=mime,
                         as_attachment=True,
                         download_name='%s.docx' %
                         doc.titleElement.getContent())

count = 0
@bp.route("/check/<doc_id>")
@user_id_required
def check(doc_id):
  logging.info("check %s", doc_id)
  user_id = request.args.get("q", "")
  logging.info("checkn doc_id = %s, user_id=%s", doc_id, user_id)
  filename = storage.get_doc_filename(get_db(), user_id, doc_id)
  if filename is None:
    file_path = None
  else:
    file_path = os.path.join(storage.data_path(current_app),
                             filename)
  if file_path is None or not os.path.isfile(file_path):
    logging.info("listen failed for file %s", file_path)
    return { "error": "Document not found" }, 400

  content = { "mtime": os.path.getmtime(file_path) }
  return flask.jsonify(content)
  
  
